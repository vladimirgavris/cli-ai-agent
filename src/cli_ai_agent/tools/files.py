from pathlib import Path

from docx import Document
from pypdf import PdfReader

from cli_ai_agent.config import KNOWLEDGE_DIR

READABLE_SUFFIXES = {".md", ".txt", ".docx", ".pdf"}
EDITABLE_SUFFIXES = {".md", ".txt", ".docx"}


def _read_lines(file_path: Path) -> list[str]:
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return [paragraph.text for paragraph in Document(str(file_path)).paragraphs]
    if suffix == ".pdf":
        reader = PdfReader(str(file_path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text.splitlines()
    return file_path.read_text(encoding="utf-8").splitlines()


def _resolve_knowledge_dir(root: str) -> Path:
    if not root:
        raise ValueError("a directory is required; use '.' for the top level")

    knowledge_root = KNOWLEDGE_DIR.resolve()
    directory = (knowledge_root / root).resolve()

    if not directory.is_relative_to(knowledge_root):
        raise ValueError("only directories inside knowledge/ are allowed")
    if not directory.is_dir():
        raise ValueError(f"directory not found: {root}")

    return directory


def list_files(root: str = ".") -> dict[str, object] | str:
    """List the folders and readable files inside one directory of knowledge/.

    The agent navigates the knowledge base one directory at a time, like `ls`:
    it starts at "." and descends using the folder names returned here.
    """
    try:
        directory = _resolve_knowledge_dir(root)
    except ValueError as error:
        return f"Cannot list directory: {error}"

    folders = sorted(
        f"{entry.name}/"
        for entry in directory.iterdir()
        if entry.is_dir() and not entry.name.startswith(".")
    )
    files = sorted(
        entry.name
        for entry in directory.iterdir()
        if entry.is_file() and entry.suffix.lower() in READABLE_SUFFIXES
    )
    relative = directory.relative_to(KNOWLEDGE_DIR.resolve())
    return {"directory": str(relative), "folders": folders, "files": files}


def _resolve_knowledge_file(name: str) -> Path:
    if not name:
        raise ValueError("a file name is required")

    knowledge_root = KNOWLEDGE_DIR.resolve()
    file_path = (knowledge_root / name).resolve()

    # FIX: was `file_path.parent != knowledge_root`, which rejected any file
    # inside a subfolder (e.g. "policies/returns.md"). Now matches the same
    # containment check used by _resolve_knowledge_dir, so subfolders work.
    if not file_path.is_relative_to(knowledge_root):
        raise ValueError("only files inside knowledge/ are allowed")

    if file_path.suffix.lower() not in READABLE_SUFFIXES:
        raise ValueError("Only .md, .txt, .docx and .pdf files are allowed")

    if not file_path.is_file():
        raise ValueError(f"knowledge file not found: {name}")

    return file_path


def read_file(name: str) -> str:
    """Read one allowed knowledge file, with every line prefixed by its number.

    The explicit "line N:" prefix gives the model exact coordinates it can
    pass to file_edit, so it never has to guess where a line is.
    """
    try:
        file_path = _resolve_knowledge_file(name)
    except ValueError as error:
        return f"Cannot read file: {error}"

    try:
        lines = _read_lines(file_path)
    except Exception:  # noqa: BLE001 - a broken file must not crash the agent
        return f"Cannot read file: {name} could not be parsed as {file_path.suffix}."

    if not lines:
        return f"{name} is empty."
    return "\n".join(
        f"line {number}: {text}".rstrip() for number, text in enumerate(lines)
    )


def _docx_replace(
    file_path: Path,
    start: int,
    end: int,
    new_paragraphs: list[str],
) -> None:
    """Splice paragraphs in place so untouched paragraphs keep their styles."""
    document = Document(str(file_path))
    paragraphs = document.paragraphs
    anchor = paragraphs[end + 1] if end + 1 < len(paragraphs) else None
    for paragraph in paragraphs[start : end + 1]:
        paragraph._element.getparent().remove(paragraph._element)
    for text in new_paragraphs:
        if anchor is None:
            document.add_paragraph(text)
        else:
            anchor.insert_paragraph_before(text)
    document.save(str(file_path))


def file_edit(
    name: str,
    start_position: int,
    end_position: int | None,
    content: str,
) -> str:
    """Replace an inclusive range of lines, or append past the end of the file.

    start_position and end_position are the 0-based line numbers that
    read_file shows. Two modes:

    * Replace: start_position points inside the file. end_position is
      required; the inclusive range [start, end] is replaced by content,
      which may hold more or fewer lines. An empty content deletes the range.
    * Append: start_position points past the last line. end_position must be
      null; the gap up to start_position is filled with empty lines and
      content starts exactly at start_position.

    For .docx files, the same line-number contract applies (one paragraph =
    one line), but the write itself is a splice via _docx_replace, so
    untouched paragraphs keep their original styles.
    """
    try:
        file_path = _resolve_knowledge_file(name)
    except ValueError as error:
        return f"Cannot edit file: {error}"

    if file_path.suffix.lower() not in EDITABLE_SUFFIXES:
        return (
            f"Cannot edit file: {file_path.suffix} files are read-only. "
            "Record the change in a .md, .txt or .docx file instead."
        )

    is_docx = file_path.suffix.lower() == ".docx"

    # FIX: previously always used file_path.read_text(), which is wrong for
    # .docx (binary zip format). Now uses the same adapter as read_file.
    try:
        lines = _read_lines(file_path)
    except Exception:  # noqa: BLE001 - a broken file must not crash the agent
        return f"Cannot edit file: {name} could not be parsed as {file_path.suffix}."

    last_line = len(lines) - 1
    if start_position < 0:
        return (
            f"Cannot edit file: start_position {start_position} is negative. "
            "Use the line numbers shown by read_file."
        )

    new_lines = content.splitlines()

    if start_position > last_line:
        # Append mode: the edit starts past the end of the file, so there is
        # no range to replace and end_position carries no meaning.
        if end_position is not None:
            return (
                f"Cannot edit file: start_position {start_position} is past the last "
                f"line ({last_line}), which appends to the file; end_position must "
                "be null in that case."
            )
        padding = start_position - len(lines)
        appended = [""] * padding + new_lines

        if is_docx:
            # Splice the new paragraphs in after the current last paragraph.
            _docx_replace(file_path, len(lines), len(lines) - 1, appended)
        else:
            lines.extend(appended)
            file_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

        new_total = start_position + len(new_lines)
        return (
            f"Edited {name}: appended {padding} empty line(s) and "
            f"{len(new_lines)} content line(s), starting at line {start_position}. "
            f"The file now has {new_total} lines. "
            "Call read_file again before another edit."
        )

    if end_position is None:
        return (
            "Cannot edit file: end_position can be null only when start_position is "
            f"past the last line ({last_line}). To replace lines, pass the last line "
            "number of the range, as shown by read_file."
        )
    if start_position > end_position or end_position > last_line:
        return (
            f"Cannot edit file: line range {start_position}-{end_position} is invalid; "
            f"{name} has lines 0-{last_line}. "
            "Call read_file to see the current line numbers."
        )

    replaced_count = end_position - start_position + 1

    if is_docx:
        _docx_replace(file_path, start_position, end_position, new_lines)
    else:
        lines[start_position : end_position + 1] = new_lines
        file_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

    new_total = len(lines) - replaced_count + len(new_lines)
    return (
        f"Edited {name}: replaced lines {start_position}-{end_position} "
        f"({replaced_count} line(s)) with {len(new_lines)} new line(s). "
        f"The file now has {new_total} lines. "
        "Line numbers may have shifted; call read_file again before another edit."
    )